import csv
from os import listdir, makedirs, path
from os.path import isfile, join
from docx import Document

STANDARD_NAME_FORMAT = 1    #1 if the name is: <last, first> it will convert it to <first last>. 0 will not change the name at all.
START_DATA = 1 #What column the data starts at, 0 indexed so START_DATA = 2 would mean that the first data column is column C
RESULTS_FOLDER = 'resultat'
RECIPE_FOLDER = 'recept'
WORD_FOLDER = 'WordFiles/' #assumes path into the folder (end with /)
req = [1, 4, 6, 6, 6, 5, 5]

def readCSVFile(file_name):
    with open(file_name, newline='') as csv_file:
        reader = csv.reader(csv_file, delimiter=';', quotechar='|')
        reader = list(reader)
        className = reader[0][0]

        print("class name", reader[0][0])
        if not path.exists(join(RECIPE_FOLDER, className)):
            makedirs(join(RECIPE_FOLDER, className))
        
        for row in reader[1:]:
            if(row[0] !=''):
                print("creating file for", row[0])
                createSingleRecipe(className, row)
            else:
                break

def createSingleRecipe(className, row):
    name = getName(row, STANDARD_NAME_FORMAT)

    isHelpNeeded = [0]*7    
    for i in range(0,7):
        if(int(row[START_DATA + i]) < req[i]):          # If the score on this part is less than required 
            isHelpNeeded[i] = 1                         # Help needed is set to one for this part
    
    if(isHelpNeeded != [0]*7):
        files = selectFiles(isHelpNeeded)
        mergeWordFiles(className, files[0],files[1], name)
    else:
        createGZCard(className, name)

def getName(row, nameFormat):
    name = row[0].strip()                               # First cell is the name of the test taker
    if(nameFormat = 1):
        name = ' '.join(name.split(', ')[::-1]).strip()     # Change namd format from "Hoogendijk, Kevin" to "Kevin Hoogendijk"
    return name

#Depending on what parts are needed for the test taker, different files are selected
def selectFiles(booleanVector):
    recipeFiles = []
    headerFiles = []
    for i in range(0,7):
        if(booleanVector[i] == 1):
            headerFiles.append(WORD_FOLDER + 'Header' + str(i+1) + '.docx')
            recipeFiles.append(WORD_FOLDER + 'Recipe' + str(i+1) + '.docx')
        else:
            headerFiles.append(WORD_FOLDER + 'HeaderDone' + str(i+1) + '.docx')
    return (recipeFiles, headerFiles)

def mergeWordFiles(className, recipeFiles, header, name):
    #can create the header by merging all the header files
    #or one can do it by manually selecting the paragraphs
    merged_document = createStandardHeader(name)
    del(merged_document.element.body[-1])
    for fileName in header:
        sub_doc = Document(fileName)
        for i, element in enumerate(sub_doc.element.body):
            if(i < len(sub_doc.element.body)):
                merged_document.element.body.append(element)
    merged_document.add_page_break()
    for i, fileName in enumerate(recipeFiles):
        sub_doc = Document(fileName)
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
    
    merged_document.save(join(RECIPE_FOLDER, className.strip(), name.strip() + '.docx'))

#A standard header is created that is needed for each document, the rest of the headers are later appended to this.
def createStandardHeader(name):
    document = Document()
    document.add_heading('Ditt matterecept ' + name + '!', 0)
    document.add_paragraph('Speciellt framtaget för att du ska kunna arbeta med rätt saker som ger'
            ' just dig så bra förutsättningar som möjligt att lyckas med matematik! '
            'Dessa uppgifter får du gärna göra under jokertiden där våra underbara Jokerlärare kan hjälpa dig. '
            'Jokrarna har även ett test på varje del som du får göra för att visa att du bemästrat den delen.')
    return(document)

#A document for the pupils that doesnt need ant recipe, a congratulation
def createGZCard(className, name):
    document = Document()
    document.add_heading(name + ' — ditt matterecept är tomt!', 0)
    document.add_paragraph('Du klarade gränserna på alla delar vilket betyder att det inte finns någon'
            ' del som behöver förbättras. Bra jobbat!')
    document.save(join(RECIPE_FOLDER, className.strip(), name + 'grattis.docx'))

#files = selectFiles([1]*7)
#mergeWordFiles("className", files[0],files[1], "test")

for f in listdir(RESULTS_FOLDER):
    readCSVFile(join(RESULTS_FOLDER, f))
